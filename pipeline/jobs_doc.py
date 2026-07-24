from __future__ import annotations

from pathlib import Path
from typing import Any


def build_jobs_document(payload: dict[str, Any], *, week_start: str, week_end: str):
    """Build a python-docx Document for the Open Positions section. Returns the Document, or None
    if python-docx is not installed (the run still succeeds; the docx is simply skipped)."""
    try:
        from docx import Document
    except ImportError:
        return None

    jobs = (payload or {}).get("jobs", [])
    doc = Document()
    doc.add_heading("Open Positions — Materials Science / Engineering", level=0)
    doc.add_paragraph(f"SoftRobotics Intelligence · {week_start} – {week_end}")
    doc.add_paragraph(
        "Current materials-science / engineering roles at robotics companies, refreshed weekly."
    )

    if not jobs:
        doc.add_paragraph("No open positions found this week.")
        return doc

    for job in jobs:
        doc.add_heading(str(job.get("title") or "Role"), level=1)
        company = str(job.get("company") or "").strip()
        location = str(job.get("location") or "").strip()
        meta = " — ".join(part for part in (company, location) if part)
        if meta:
            doc.add_paragraph(meta)
        description = str(job.get("description") or "").strip()
        if description:
            doc.add_paragraph(description)
        url = str(job.get("url") or "").strip()
        if url:
            doc.add_paragraph(url)

    return doc


def write_jobs_doc(
    payload: dict[str, Any],
    *,
    week_start: str,
    week_end: str,
    output_dir: str = "output",
) -> Path | None:
    """Write the weekly Open Positions Word document to output_dir. Returns the written path, or
    None if python-docx is unavailable."""
    doc = build_jobs_document(payload, week_start=week_start, week_end=week_end)
    if doc is None:
        return None
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    path = out / f"SoftRobotics_Jobs_{week_end}.docx"
    doc.save(str(path))
    return path
